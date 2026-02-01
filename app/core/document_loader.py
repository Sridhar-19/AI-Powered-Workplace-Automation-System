"""
Document loaders for various file formats.
"""
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from langchain.schema import Document
from langchain_community.document_loaders import UnstructuredFileLoader

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Universal document loader supporting multiple file formats.
    """

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']

    def load(self, file_path: str) -> List[Document]:
        """
        Load a document from file path.
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of Document objects
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        try:
            if extension == '.pdf':
                return self._load_pdf(file_path)
            elif extension == '.docx':
                return self._load_docx(file_path)
            elif extension in ['.txt', '.md']:
                return self._load_text(file_path)
            else:
                # Fallback to unstructured loader
                return self._load_unstructured(file_path)
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            raise

    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load PDF document."""
        try:
            documents = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Extract metadata
                metadata = {
                    'source': file_path,
                    'file_type': 'pdf',
                    'num_pages': num_pages,
                }
                
                # Add PDF metadata if available
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['creation_date'] = pdf_reader.metadata.get('/CreationDate', '')
                
                # Extract text from each page
                full_text = []
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        full_text.append(text)
                        
                        # Create document for each page
                        page_metadata = metadata.copy()
                        page_metadata['page'] = page_num + 1
                        
                        documents.append(Document(
                            page_content=text,
                            metadata=page_metadata
                        ))
                
                logger.info(f"Loaded PDF with {num_pages} pages from {file_path}")
                return documents
                
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise

    def _load_docx(self, file_path: str) -> List[Document]:
        """Load DOCX document."""
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                'source': file_path,
                'file_type': 'docx',
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'created': str(core_properties.created) if core_properties.created else '',
                'modified': str(core_properties.modified) if core_properties.modified else '',
            }
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Combine into single document
            full_text = '\n\n'.join(paragraphs)
            
            # Also extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            if tables_text:
                full_text += '\n\nTables:\n' + '\n'.join(tables_text)
            
            logger.info(f"Loaded DOCX with {len(paragraphs)} paragraphs from {file_path}")
            
            return [Document(page_content=full_text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise

    def _load_text(self, file_path: str) -> List[Document]:
        """Load plain text or markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            path = Path(file_path)
            metadata = {
                'source': file_path,
                'file_type': path.suffix[1:],  # Remove the dot
            }
            
            logger.info(f"Loaded text file from {file_path}")
            
            return [Document(page_content=text, metadata=metadata)]
            
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise

    def _load_unstructured(self, file_path: str) -> List[Document]:
        """
        Load document using unstructured library for complex formats.
        """
        try:
            loader = UnstructuredFileLoader(file_path)
            documents = loader.load()
            
            logger.info(f"Loaded document using unstructured from {file_path}")
            
            return documents
            
        except Exception as e:
            logger.error(f"Error loading with unstructured {file_path}: {e}")
            raise

    def load_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        file_type: str
    ) -> List[Document]:
        """
        Load document from bytes (useful for uploaded files).
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            file_type: File extension (e.g., 'pdf', 'docx')
            
        Returns:
            List of Document objects
        """
        import tempfile
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=f'.{file_type}'
        ) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        try:
            # Load from temporary file
            documents = self.load(tmp_path)
            
            # Update metadata with original filename
            for doc in documents:
                doc.metadata['original_filename'] = filename
            
            return documents
            
        finally:
            # Clean up temporary file
            Path(tmp_path).unlink(missing_ok=True)


def create_document_loader() -> DocumentLoader:
    """Factory function to create document loader."""
    return DocumentLoader()
